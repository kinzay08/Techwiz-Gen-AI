


import os
import requests
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_pymongo import PyMongo
from flask_bcrypt import Bcrypt
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_dance.contrib.google import make_google_blueprint, google
from bson import ObjectId
from datetime import datetime, timedelta
from dotenv import load_dotenv
from dateutil.parser import parse
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import tempfile
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import re

# Import AI services (assumed to be in services/ directory)
from services.whisper_service import transcribe_audio
from services.gemini_service import summarize_with_gemini
from services.grok_service import summarize_with_groq, get_available_models

# Load environment variables
load_dotenv()
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

# Flask app setup
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "supersecretkey")
app.config["MONGO_URI"] = os.getenv("MONGO_URI", "mongodb://localhost:27017/flask_auth")
mongo = PyMongo(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = "login"

# Google OAuth Config
app.config["GOOGLE_OAUTH_CLIENT_ID"] = os.getenv("GOOGLE_OAUTH_CLIENT_ID")
app.config["GOOGLE_OAUTH_CLIENT_SECRET"] = os.getenv("GOOGLE_OAUTH_CLIENT_SECRET")
google_bp = make_google_blueprint(
    scope=["openid", "https://www.googleapis.com/auth/userinfo.profile", "https://www.googleapis.com/auth/userinfo.email"],
    redirect_to="google_login"
)
app.register_blueprint(google_bp, url_prefix="/login")
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", 587))
EMAIL_USER = os.getenv("EMAIL_USER")  # Your email
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")  # Your email password or app password
FROM_EMAIL = os.getenv("FROM_EMAIL", EMAIL_USER)

# OpenRouter API configuration
API_KEY = os.getenv("OPENROUTER_API_KEY")
BASE_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "nvidia/nemotron-nano-9b-v2:free"
SYSTEM_PROMPT = """You are an AI meeting assistant. Your role is to analyze raw transcripts and generate structured, concise notes.

Always format output in this structure:

**Executive Summary:**
Brief overview of the meeting purpose and main outcomes.

**Key Discussion Points:**
• Main topics discussed
• Important details and context

**Action Items:**
• Specific tasks with responsible parties (if mentioned)
• Deadlines (if mentioned)

**Decisions Made:**
• Clear decisions reached during the meeting
• Next steps agreed upon

Please be concise but comprehensive in your analysis."""

if not API_KEY:
    print("Warning: OPENROUTER_API_KEY not found in environment variables")

# Helper functions for export
def format_notes_for_docx(notes):
    """Parse structured notes for Word export."""
    formatted = []
    lines = notes.split('\n') if notes else []
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('**') and line.endswith('**'):
            current_section = line.strip('*:').strip()
            formatted.append(('Heading 2', current_section))
        elif line.startswith('•'):
            formatted.append(('List Bullet', line.lstrip('• ').strip()))
        else:
            formatted.append(('Normal', line))
    
    return formatted

def create_word_document(note):
    """Create a Word document from a note."""
    doc = Document()
    
    doc.add_heading(note.get('title', 'Untitled Note'), level=1)
    doc.add_paragraph(f"Created: {note.get('created_at', 'Unknown')}", style='Normal')
    doc.add_paragraph(f"AI Provider: {note.get('ai_provider', 'Unknown')}", style='Normal')
    doc.add_paragraph(f"Updated: {note.get('updated_at', 'Unknown')}", style='Normal')
    
    if note.get('transcript'):
        doc.add_heading('Transcript', level=2)
        doc.add_paragraph(note['transcript'], style='Normal')
    
    if note.get('notes'):
        doc.add_heading('Notes', level=2)
        formatted_notes = format_notes_for_docx(note.get('notes', ''))
        for style, content in formatted_notes:
            p = doc.add_paragraph(content, style=style)
            if style == 'Normal' or style == 'List Bullet':
                for run in p.runs:
                    run.font.size = Pt(11)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_document(note):
    """Create a PDF document using ReportLab from a note."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=1*inch, rightMargin=1*inch, topMargin=1*inch, bottomMargin=1*inch)
    styles = getSampleStyleSheet()
    
    # Define custom styles
    title_style = styles['Title']
    heading_style = styles['Heading2']
    normal_style = ParagraphStyle(name='Normal', fontSize=11, leading=14, spaceAfter=12)
    bullet_style = ParagraphStyle(name='Bullet', fontSize=11, leading=14, leftIndent=20, bulletIndent=10, spaceAfter=12)
    
    story = []
    
    # Title
    story.append(Paragraph(note.get('title', 'Untitled Note'), title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Metadata
    story.append(Paragraph('Metadata', heading_style))
    metadata = [
        f"<b>Created:</b> {note.get('created_at', 'Unknown')}",
        f"<b>Updated:</b> {note.get('updated_at', 'Unknown')}",
        f"<b>AI Provider:</b> {note.get('ai_provider', 'Unknown')}"
    ]
    for item in metadata:
        story.append(Paragraph(item, normal_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Transcript
    if note.get('transcript'):
        story.append(Paragraph('Transcript', heading_style))
        transcript_lines = note['transcript'].split('\n')
        for line in transcript_lines:
            if line.strip():
                story.append(Paragraph(line.strip(), normal_style))
        story.append(Spacer(1, 0.2*inch))
    
    # Notes
    if note.get('notes'):
        story.append(Paragraph('Notes', heading_style))
        lines = note['notes'].split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('**') and line.endswith('**'):
                story.append(Paragraph(line.strip('*:').strip(), heading_style))
            elif line.startswith('•'):
                story.append(Paragraph(f"• {line.lstrip('• ').strip()}", bullet_style))
            else:
                story.append(Paragraph(line, normal_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer

def summarize_with_openai(transcript: str) -> str:
    """Use OpenRouter API to summarize transcript into structured notes."""
    if not transcript.strip():
        return "Error: No transcript provided to summarize."
    if not API_KEY:
        return "Error: OpenRouter API key not configured."
    
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "http://localhost:5000",
        "X-Title": "Meeting Notes AI"
    }
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Please analyze this meeting transcript and create structured notes:\n\n{transcript}"}
        ],
        "max_tokens": 1000,
        "temperature": 0.7,
        "top_p": 0.8,
        "frequency_penalty": 0.1,
        "presence_penalty": 0.1
    }
    
    try:
        print(f"🚀 Sending request to OpenRouter API with model: {MODEL_NAME}")
        response = requests.post(BASE_URL, headers=headers, json=payload, timeout=30)
        print(f"📡 Response Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            if 'choices' in data and len(data['choices']) > 0:
                message = data['choices'][0].get('message', {})
                content = message.get('content', '')
                if content:
                    print(f"✅ OpenRouter response received: {len(content)} characters")
                    return content.strip()
                return "Error: Empty response from OpenRouter"
            return "Error: Invalid response format from OpenRouter"
        elif response.status_code == 401:
            return "Error: Invalid API key for OpenRouter"
        elif response.status_code == 402:
            return "Error: Insufficient credits in OpenRouter account"
        elif response.status_code == 429:
            return "Error: Rate limit exceeded for OpenRouter API"
        else:
            error_text = response.text
            print(f"❌ API Error {response.status_code}: {error_text}")
            return f"Error {response.status_code}: Failed to process with OpenRouter"
    except requests.exceptions.Timeout:
        print("❌ Request timeout")
        return "Error: Request timed out while processing with OpenRouter"
    except requests.exceptions.RequestException as e:
        print(f"❌ Request exception: {str(e)}")
        return f"Error: Request failed - {str(e)}"
    except Exception as e:
        print(f"❌ Unexpected error: {str(e)}")
        return f"Error: Unexpected error - {str(e)}"

def get_openai_model_status():
    """Check if OpenRouter service is available."""
    return {
        'available': bool(API_KEY),
        'model': MODEL_NAME,
        'provider': 'openrouter'
    }

def get_ai_provider_status():
    """Check which AI providers are available."""
    status = {
        'openai': {
            'available': bool(os.getenv("OPENROUTER_API_KEY")),
            'name': 'OpenRouter (NVIDIA Nemotron)',
            'model': MODEL_NAME
        },
        'groq': {
            'available': bool(os.getenv("GROQ_API_KEY")),
            'name': 'Groq',
            'models': get_available_models() if os.getenv("GROQ_API_KEY") else []
        },
        'gemini': {
            'available': bool(os.getenv("GOOGLE_API_KEY")),
            'name': 'Google Gemini'
        }
    }
    return status

def summarize_with_ai(transcript, provider="auto", groq_model="llama-3.3-70b-versatile"):
    """Summarize transcript using available AI provider."""
    print(f"🔧 summarize_with_ai called with provider='{provider}', model='{groq_model}'")
    if not transcript or not transcript.strip():
        print("❌ No transcript provided")
        return "Error: No transcript provided.", "none"
    
    providers_status = get_ai_provider_status()
    print(f"📊 Providers status: {providers_status}")
    
    if provider == "auto":
        if providers_status['openai']['available']:
            provider = "openai"
            print("🔄 Auto-selected OpenRouter provider")
        elif providers_status['groq']['available']:
            provider = "groq"
            print("🔄 Auto-selected Groq provider")
        elif providers_status['gemini']['available']:
            provider = "gemini"
            print("🔄 Auto-selected Gemini provider")
        else:
            error_msg = "Error: No AI providers are available. Please configure API keys."
            print(f"❌ {error_msg}")
            return error_msg, "none"
    
    print(f"🎯 Using provider: {provider}")
    
    if provider == "openai" and providers_status['openai']['available']:
        try:
            print("🚀 Calling OpenRouter...")
            start_time = datetime.now()
            summary = summarize_with_openai(transcript)
            processing_time = (datetime.now() - start_time).total_seconds()
            log_user_activity(
                current_user.id,
                "ai_summary",
                "Generated summary using OpenRouter",
                {
                    "provider": "openai",
                    "model": MODEL_NAME,
                    "transcript_length": len(transcript),
                    "processing_time_seconds": processing_time,
                    "success": not summary.startswith("Error")
                }
            )
            if summary and not summary.startswith("Error"):
                print("✅ OpenRouter successful")
                update_user_stats(current_user.id, "ai_summary", 1, {"processing_time": processing_time})
                return summary, "openai"
            print(f"❌ OpenRouter failed: {summary}")
            if providers_status['groq']['available']:
                try:
                    print("🔄 Falling back to Groq...")
                    summary = summarize_with_groq(transcript, groq_model)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        f"Generated summary using Groq (fallback from OpenRouter) - {groq_model}",
                        {
                            "provider": "groq",
                            "model": groq_model,
                            "fallback_from": "openai",
                            "transcript_length": len(transcript),
                            "success": not summary.startswith("Error")
                        }
                    )
                    return summary, "groq (fallback)"
                except Exception as e:
                    print(f"❌ Groq fallback failed: {e}")
            if providers_status['gemini']['available']:
                try:
                    print("🔄 Falling back to Gemini...")
                    summary = summarize_with_gemini(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using Gemini (fallback from OpenRouter)",
                        {
                            "provider": "gemini",
                            "fallback_from": "openai",
                            "transcript_length": len(transcript),
                            "success": not summary.startswith("Error")
                        }
                    )
                    return summary, "gemini (fallback)"
                except Exception as e:
                    print(f"❌ Gemini fallback failed: {e}")
            return summary, "none"
        except Exception as e:
            error_msg = f"OpenRouter failed: {str(e)}"
            print(f"❌ Exception with OpenRouter: {error_msg}")
            log_user_activity(
                current_user.id,
                "error",
                error_msg,
                {"error_type": "ai_summary", "provider": "openai"}
            )
            if providers_status['groq']['available']:
                try:
                    print("🔄 Exception fallback to Groq...")
                    summary = summarize_with_groq(transcript, groq_model)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        f"Generated summary using Groq (fallback from OpenRouter exception) - {groq_model}",
                        {
                            "provider": "groq",
                            "model": groq_model,
                            "fallback_from": "openai",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "groq (fallback)"
                except Exception as e:
                    print(f"❌ Groq fallback failed: {e}")
            if providers_status['gemini']['available']:
                try:
                    print("🔄 Exception fallback to Gemini...")
                    summary = summarize_with_gemini(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using Gemini (fallback from OpenRouter exception)",
                        {
                            "provider": "gemini",
                            "fallback_from": "openai",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "gemini (fallback)"
                except Exception as e:
                    print(f"❌ Gemini fallback failed: {e}")
            return error_msg, "none"
    
    elif provider == "groq" and providers_status['groq']['available']:
        try:
            print(f"🚀 Calling Groq with model: {groq_model}")
            start_time = datetime.now()
            summary = summarize_with_groq(transcript, groq_model)
            processing_time = (datetime.now() - start_time).total_seconds()
            log_user_activity(
                current_user.id,
                "ai_summary",
                f"Generated summary using Groq ({groq_model})",
                {
                    "provider": "groq",
                    "model": groq_model,
                    "transcript_length": len(transcript),
                    "processing_time_seconds": processing_time,
                    "success": not summary.startswith("Error")
                }
            )
            if summary and not summary.startswith("Error"):
                print("✅ Groq successful")
                update_user_stats(current_user.id, "ai_summary", 1, {"processing_time": processing_time})
                return summary, "groq"
            print(f"❌ Groq failed: {summary}")
            if providers_status['openai']['available']:
                try:
                    print("🔄 Falling back to OpenRouter...")
                    summary = summarize_with_openai(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using OpenRouter (fallback from Groq)",
                        {
                            "provider": "openai",
                            "fallback_from": "groq",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "openai (fallback)"
                except Exception as e:
                    print(f"❌ OpenRouter fallback failed: {e}")
            if providers_status['gemini']['available']:
                try:
                    print("🔄 Falling back to Gemini...")
                    summary = summarize_with_gemini(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using Gemini (fallback from Groq)",
                        {
                            "provider": "gemini",
                            "fallback_from": "groq",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "gemini (fallback)"
                except Exception as e:
                    print(f"❌ Gemini fallback failed: {e}")
            return summary, "none"
        except Exception as e:
            error_msg = f"Groq failed: {str(e)}"
            print(f"❌ Exception with Groq: {error_msg}")
            log_user_activity(
                current_user.id,
                "error",
                error_msg,
                {"error_type": "ai_summary", "provider": "groq"}
            )
            if providers_status['openai']['available']:
                try:
                    print("🔄 Exception fallback to OpenRouter...")
                    summary = summarize_with_openai(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using OpenRouter (fallback from Groq exception)",
                        {
                            "provider": "openai",
                            "fallback_from": "groq",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "openai (fallback)"
                except Exception as e:
                    print(f"❌ OpenRouter fallback failed: {e}")
            if providers_status['gemini']['available']:
                try:
                    print("🔄 Exception fallback to Gemini...")
                    summary = summarize_with_gemini(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        f"Generated summary using Gemini (fallback from Groq exception)",
                        {
                            "provider": "gemini",
                            "fallback_from": "groq",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "gemini (fallback)"
                except Exception as e:
                    print(f"❌ Gemini fallback failed: {e}")
            return error_msg, "none"
    
    elif provider == "gemini" and providers_status['gemini']['available']:
        try:
            print("🚀 Calling Gemini...")
            start_time = datetime.now()
            summary = summarize_with_gemini(transcript)
            processing_time = (datetime.now() - start_time).total_seconds()
            log_user_activity(
                current_user.id,
                "ai_summary",
                "Generated summary using Gemini",
                {
                    "provider": "gemini",
                    "transcript_length": len(transcript),
                    "processing_time_seconds": processing_time,
                    "success": not summary.startswith("Error")
                }
            )
            if summary and not summary.startswith("Error"):
                print("✅ Gemini successful")
                update_user_stats(current_user.id, "ai_summary", 1, {"processing_time": processing_time})
                return summary, "gemini"
            print(f"❌ Gemini failed: {summary}")
            if providers_status['openai']['available']:
                try:
                    print("🔄 Falling back to OpenRouter...")
                    summary = summarize_with_openai(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using OpenRouter (fallback from Gemini)",
                        {
                            "provider": "openai",
                            "fallback_from": "gemini",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "openai (fallback)"
                except Exception as e:
                    print(f"❌ OpenRouter fallback failed: {e}")
            if providers_status['groq']['available']:
                try:
                    print("🔄 Falling back to Groq...")
                    summary = summarize_with_groq(transcript, groq_model)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        f"Generated summary using Groq (fallback from Gemini) - {groq_model}",
                        {
                            "provider": "groq",
                            "model": groq_model,
                            "fallback_from": "gemini",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "groq (fallback)"
                except Exception as e:
                    print(f"❌ Groq fallback failed: {e}")
            return summary, "none"
        except Exception as e:
            error_msg = f"Gemini failed: {str(e)}"
            print(f"❌ Exception with Gemini: {error_msg}")
            log_user_activity(
                current_user.id,
                "error",
                error_msg,
                {"error_type": "ai_summary", "provider": "gemini"}
            )
            if providers_status['openai']['available']:
                try:
                    print("🔄 Exception fallback to OpenRouter...")
                    summary = summarize_with_openai(transcript)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        "Generated summary using OpenRouter (fallback from Gemini exception)",
                        {
                            "provider": "openai",
                            "fallback_from": "gemini",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "openai (fallback)"
                except Exception as e:
                    print(f"❌ OpenRouter fallback failed: {e}")
            if providers_status['groq']['available']:
                try:
                    print("🔄 Exception fallback to Groq...")
                    summary = summarize_with_groq(transcript, groq_model)
                    log_user_activity(
                        current_user.id,
                        "ai_summary",
                        f"Generated summary using Groq (fallback from Gemini exception) - {groq_model}",
                        {
                            "provider": "groq",
                            "model": groq_model,
                            "fallback_from": "gemini",
                            "transcript_length": len(transcript)
                        }
                    )
                    return summary, "groq (fallback)"
                except Exception as e:
                    print(f"❌ Groq fallback failed: {e}")
            return error_msg, "none"
    
    else:
        error_msg = f"Error: Provider '{provider}' is not available or not configured."
        print(f"❌ {error_msg}")
        return error_msg, "none"

def log_user_activity(user_id, activity_type, details=None, metadata=None):
    """Log user activity to the database."""
    try:
        if not ObjectId.is_valid(str(user_id)):
            print(f"❌ Invalid user_id: {user_id}")
            return
        activity_data = {
            "user_id": ObjectId(str(user_id)),
            "activity_type": activity_type,
            "details": details or "No details provided",
            "metadata": metadata or {},
            "timestamp": datetime.now(),
            "ip_address": request.environ.get('HTTP_X_REAL_IP', request.remote_addr),
            "user_agent": request.headers.get('User-Agent')
        }
        mongo.db.user_activities.insert_one(activity_data)
        print(f"📊 Activity logged: {activity_type} for user {user_id}")
    except Exception as e:
        print(f"❌ Failed to log activity: {e}")

def update_user_stats(user_id, stat_type, increment=1, additional_data=None):
    """Update user statistics."""
    try:
        if not ObjectId.is_valid(str(user_id)):
            print(f"❌ Invalid user_id: {user_id}")
            return
        today = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        daily_stats = {
            "user_id": ObjectId(str(user_id)),
            "date": today,
            f"{stat_type}_count": increment,
            "last_updated": datetime.now()
        }
        if additional_data:
            daily_stats.update({k: v for k, v in additional_data.items() if isinstance(v, (int, float))})
        mongo.db.daily_stats.update_one(
            {"user_id": ObjectId(str(user_id)), "date": today},
            {"$inc": {f"{stat_type}_count": increment},
             "$set": {"last_updated": datetime.now()},
             "$setOnInsert": daily_stats},
            upsert=True
        )
        user_stats = {
            f"total_{stat_type}": increment,
            f"last_{stat_type}_date": datetime.now()
        }
        if additional_data:
            for key, value in additional_data.items():
                if isinstance(value, (int, float)):
                    user_stats[f"total_{key}"] = value
        mongo.db.user_stats.update_one(
            {"user_id": ObjectId(str(user_id))},
            {"$inc": {f"total_{stat_type}": increment},
             "$set": {f"last_{stat_type}_date": datetime.now()},
             "$setOnInsert": {"user_id": ObjectId(str(user_id)), "created_at": datetime.now()}},
            upsert=True
        )
        print(f"📈 Stats updated: {stat_type} +{increment} for user {user_id}")
    except Exception as e:
        print(f"❌ Failed to update stats: {e}")

def get_user_statistics(user_id, days=30):
    """Get comprehensive user statistics."""
    try:
        if not ObjectId.is_valid(str(user_id)):
            print(f"❌ Invalid user_id: {user_id}")
            return {}
        stats = {
            "overview": {},
            "daily_activity": [],
            "ai_provider_usage": [],
            "activity_timeline": [],
            "file_processing": {},
            "peak_usage": []
        }
        user_stats = mongo.db.user_stats.find_one({"user_id": ObjectId(str(user_id))})
        if user_stats:
            created_at = user_stats.get("created_at")
            if isinstance(created_at, str):
                created_at = parse(created_at)
            last_activity = user_stats.get("last_login_date")
            if isinstance(last_activity, str):
                last_activity = parse(last_activity)
            stats["overview"] = {
                "total_transcriptions": user_stats.get("total_transcription", 0),
                "total_ai_summaries": user_stats.get("total_ai_summary", 0),
                "total_notes_saved": user_stats.get("total_save_notes", 0),
                "total_logins": user_stats.get("total_login", 0),
                "account_created": created_at or datetime.now(),
                "last_activity": last_activity or datetime.now()
            }
        start_date = datetime.now() - timedelta(days=days)
        daily_data = list(mongo.db.daily_stats.find({
            "user_id": ObjectId(str(user_id)),
            "date": {"$gte": start_date}
        }).sort("date", 1))
        stats["daily_activity"] = [
            {
                "date": day["date"].strftime("%Y-%m-%d") if isinstance(day["date"], datetime) else day["date"],
                "transcription_count": day.get("transcription_count", 0),
                "ai_summary_count": day.get("ai_summary_count", 0),
                "save_notes_count": day.get("save_notes_count", 0),
                "login_count": day.get("login_count", 0)
            } for day in daily_data
        ]
        pipeline = [
            {"$match": {"user_id": ObjectId(str(user_id)), "metadata.provider": {"$exists": True}}},
            {"$group": {
                "_id": {"$toLower": "$metadata.provider"},
                "count": {"$sum": 1},
                "last_used": {"$max": "$timestamp"}
            }},
            {"$sort": {"count": -1}}
        ]
        ai_usage = list(mongo.db.user_activities.aggregate(pipeline))
        stats["ai_provider_usage"] = [
            {
                "_id": item["_id"],
                "count": item["count"],
                "last_used": item["last_used"].strftime("%Y-%m-%d %H:%M:%S") if isinstance(item["last_used"], datetime) else item["last_used"]
            } for item in ai_usage
        ]
        recent_activities = list(mongo.db.user_activities.find(
            {"user_id": ObjectId(str(user_id))}
        ).sort("timestamp", -1).limit(50))
        stats["activity_timeline"] = [
            {
                "activity_type": activity["activity_type"],
                "details": activity.get("details", "No details"),
                "metadata": activity.get("metadata", {}),
                "timestamp": activity["timestamp"].strftime("%Y-%m-%d %H:%M:%S") if isinstance(activity["timestamp"], datetime) else activity["timestamp"]
            } for activity in recent_activities
        ]
        processing_pipeline = [
            {"$match": {"user_id": ObjectId(str(user_id)), "activity_type": "transcription"}},
            {"$group": {
                "_id": None,
                "total_files": {"$sum": 1},
                "total_duration": {"$sum": {"$ifNull": ["$metadata.duration_seconds", 0]}},
                "avg_file_size": {"$avg": {"$ifNull": ["$metadata.file_size_mb", 0]}},
                "most_common_format": {"$push": {"$ifNull": ["$metadata.file_format", "unknown"]}}
            }}
        ]
        file_stats = list(mongo.db.user_activities.aggregate(processing_pipeline))
        if file_stats:
            stats["file_processing"] = {
                "total_files": file_stats[0].get("total_files", 0),
                "total_duration": file_stats[0].get("total_duration", 0),
                "avg_file_size": file_stats[0].get("avg_file_size", 0),
                "most_common_format": max(set(file_stats[0].get("most_common_format", ["unknown"])), key=lambda x: file_stats[0].get("most_common_format", ["unknown"]).count(x))
            }
        peak_pipeline = [
            {"$match": {"user_id": ObjectId(str(user_id))}},
            {"$group": {
                "_id": {"$hour": "$timestamp"},
                "activity_count": {"$sum": 1}
            }},
            {"$sort": {"activity_count": -1}},
            {"$limit": 3}
        ]
        peak_hours = list(mongo.db.user_activities.aggregate(peak_pipeline))
        stats["peak_usage"] = [
            {
                "_id": str(item["_id"]),
                "activity_count": item["activity_count"]
            } for item in peak_hours
        ]
        return stats
    except Exception as e:
        print(f"❌ Error getting user statistics: {e}")
        log_user_activity(user_id, "error", f"Failed to fetch statistics: {str(e)}", {"error_type": "statistics"})
        return {}

# User Model
class User(UserMixin):
    def __init__(self, user_data):
        self.id = str(user_data["_id"])
        self.username = user_data.get("username", "")
        self.email = user_data["email"]
        self.fullname = user_data.get("fullname", "")

@login_manager.user_loader
def load_user(user_id):
    if not ObjectId.is_valid(user_id):
        return None
    user_data = mongo.db.users.find_one({"_id": ObjectId(user_id)})
    return User(user_data) if user_data else None

@app.template_filter('moment')
def moment_filter(value=None, fmt="%b %d, %Y"):
    if not value:
        return datetime.now().strftime(fmt)
    try:
        if isinstance(value, str):
            value = parse(value)  # Convert string to datetime
        return value.strftime(fmt)  # Format datetime
    except Exception:
        return str(value)  # If anything goes wrong, return the value as is

# Routes
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/overview")
def overview():
    return render_template("overview.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        fullname = request.form.get("fullname", "").strip()
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm_password = request.form.get("confirm_password", "")
        phone = request.form.get("phone", "")
        dob = request.form.get("dob", "")
        if not all([fullname, username, email, password, confirm_password]):
            flash("All required fields must be filled!", "danger")
            return redirect(url_for("register"))
        if password != confirm_password:
            flash("Passwords do not match!", "danger")
            return redirect(url_for("register"))
        if mongo.db.users.find_one({"email": email}):
            flash("Email already exists!", "danger")
            return redirect(url_for("register"))
        hashed_pw = bcrypt.generate_password_hash(password).decode("utf-8")
        user_id = mongo.db.users.insert_one({
            "fullname": fullname,
            "username": username,
            "email": email,
            "password": hashed_pw,
            "phone": phone,
            "dob": dob,
            "oauth": False,
            "created_at": datetime.now().isoformat()
        }).inserted_id
        user = mongo.db.users.find_one({"_id": user_id})
        login_user(User(user))
        log_user_activity(str(user_id), "register", "User account created", {"method": "email"})
        update_user_stats(str(user_id), "register")
        return redirect(url_for("overview"))
    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        if not email or not password:
            flash("Email and password are required!", "danger")
            return redirect(url_for("login"))
        user = mongo.db.users.find_one({"email": email, "oauth": False})
        if user and bcrypt.check_password_hash(user["password"], password):
            login_user(User(user))
            log_user_activity(str(user["_id"]), "login", "User logged in", {"method": "email"})
            update_user_stats(str(user["_id"]), "login")
            return redirect(url_for("overview"))
        flash("Invalid email or password", "danger")
    return render_template("login.html")

@app.route("/google_login")
def google_login():
    if not google.authorized:
        return redirect(url_for("google.login"))
    resp = google.get("/oauth2/v2/userinfo")
    if not resp.ok:
        flash("Failed to fetch user info from Google.", "danger")
        return redirect(url_for("login"))
    info = resp.json()
    email = info["email"]
    user = mongo.db.users.find_one({"email": email})
    if not user:
        user_id = mongo.db.users.insert_one({
            "username": info.get("name", ""),
            "fullname": info.get("name", ""),
            "email": email,
            "password": None,
            "oauth": True,
            "created_at": datetime.now().isoformat()
        }).inserted_id
        user = mongo.db.users.find_one({"_id": user_id})
        log_user_activity(str(user_id), "register", "User account created via Google OAuth", {"method": "google"})
        update_user_stats(str(user_id), "register")
    login_user(User(user))
    log_user_activity(str(user["_id"]), "login", "User logged in via Google OAuth", {"method": "google"})
    update_user_stats(str(user["_id"]), "login")
    return redirect(url_for("overview"))

@app.route("/dashboard", methods=["GET", "POST"])
@login_required
def dashboard():
    transcript = None
    notes = None
    provider_used = None
    audio_path = None  # Initialize here!
    ai_providers = get_ai_provider_status()
    
    if request.method == "POST":
        if "audio_file" not in request.files:
            flash("No file uploaded!", "danger")
            return redirect(url_for("dashboard"))
        file = request.files["audio_file"]
        if file.filename == "":
            flash("No file selected!", "danger")
            return redirect(url_for("dashboard"))
        ai_provider = request.form.get("ai_provider", "auto")
        groq_model = request.form.get("groq_model", "llama-3.3-70b-versatile")
        print(f"🎯 Processing with provider: {ai_provider}, model: {groq_model}")
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp:
                file.save(temp.name)
                audio_path = temp.name
            print(f"🎵 Audio saved to: {audio_path}")
            file_size_mb = os.path.getsize(audio_path) / (1024 * 1024)
            file_format = file.filename.split('.')[-1].lower() if '.' in file.filename else 'unknown'
            print("🔄 Starting transcription...")
            start_time = datetime.now()
            transcript = transcribe_audio(audio_path, language="en")
            transcription_time = (datetime.now() - start_time).total_seconds()
            
            if transcript:
                print(f"✅ Transcription successful: {len(transcript)} characters")
                print(f"📝 Transcript preview: {transcript[:100]}...")
                log_user_activity(
                    current_user.id,
                    "transcription",
                    f"Audio file transcribed: {file.filename}",
                    {
                        "filename": file.filename,
                        "file_size_mb": file_size_mb,
                        "file_format": file_format,
                        "transcript_length": len(transcript),
                        "duration_seconds": transcription_time,
                        "success": True
                    }
                )
                update_user_stats(current_user.id, "transcription", 1, {"file_size": file_size_mb})
                print(f"🤖 Starting AI summarization with {ai_provider}...")
                result = summarize_with_ai(transcript, ai_provider, groq_model)
                if isinstance(result, tuple) and len(result) == 2:
                    notes, provider_used = result
                    print(f"✅ AI processing successful with {provider_used}")
                    print(f"📄 Notes preview: {notes[:200]}...")
                    if notes and not notes.startswith("Error"):
                        flash(f"Audio processed successfully using {provider_used}!", "success")
                    else:
                        flash(f"AI Processing failed: {notes}", "danger")
                        print(f"❌ AI failed with error: {notes}")
                else:
                    notes = "Error: Invalid AI response format"
                    provider_used = "none"
                    flash("AI Processing failed: Invalid response format", "danger")
                    print(f"❌ Invalid AI response format: {result}")
            else:
                flash("Failed to transcribe audio", "danger")
                print("❌ Transcription failed")
                log_user_activity(
                    current_user.id,
                    "error",
                    "Transcription failed",
                    {"error_type": "transcription", "filename": file.filename}
                )
        except Exception as e:
            error_msg = str(e)
            print(f"💥 Exception in dashboard: {error_msg}")
            log_user_activity(
                current_user.id,
                "error",
                f"Audio processing failed: {error_msg}",
                {"error_type": "processing", "filename": file.filename if 'file' in locals() else "unknown"}
            )
            flash(f"Error processing audio: {error_msg}", "danger")
        finally:
            # Clean up temp file if it exists
            if audio_path and os.path.exists(audio_path):
                os.remove(audio_path)
                print(f"🗑️ Cleaned up temp file: {audio_path}")
    
    name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
    quick_stats = get_user_statistics(current_user.id, days=7)
    return render_template("dashboard.html",
                         name=name,
                         transcript=transcript,
                         notes=notes,
                         provider_used=provider_used,
                         ai_providers=ai_providers,
                         quick_stats=quick_stats)
@app.route("/save_notes", methods=["POST"])
@login_required
def save_notes():
    try:
        transcript = request.form.get("transcript", "")
        notes = request.form.get("notes", "")
        title = request.form.get("title", f"Meeting Notes - {datetime.now().strftime('%B %d, %Y')}")
        provider_used = request.form.get("provider_used", "unknown")
        if not notes:
            flash("No notes to save!", "danger")
            return redirect(url_for("dashboard"))
        note_data = {
            "user_id": ObjectId(current_user.id),
            "title": title,
            "transcript": transcript,
            "notes": notes,
            "ai_provider": provider_used,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        result = mongo.db.saved_notes.insert_one(note_data)
        if result.inserted_id:
            log_user_activity(
                current_user.id,
                "save_notes",
                f"Notes saved: {title}",
                {
                    "note_id": str(result.inserted_id),
                    "ai_provider": provider_used,
                    "transcript_length": len(transcript),
                    "notes_length": len(notes)
                }
            )
            update_user_stats(current_user.id, "save_notes")
            flash("Notes saved successfully!", "success")
        else:
            flash("Failed to save notes!", "danger")
    except Exception as e:
        log_user_activity(
            current_user.id,
            "error",
            f"Failed to save notes: {str(e)}",
            {"error_type": "save_notes"}
        )
        flash(f"Error saving notes: {str(e)}", "danger")
    return redirect(url_for("dashboard"))

@app.route("/statistics")
@login_required
def statistics():
    try:
        stats = get_user_statistics(current_user.id, days=30)
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        if not stats:
            flash("No statistics available at this time.", "warning")
            return redirect(url_for("dashboard"))
        return render_template("statistics.html", stats=stats, name=name)
    except Exception as e:
        flash(f"Error loading statistics: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to load statistics: {str(e)}", {"error_type": "statistics"})
        return redirect(url_for("dashboard"))

ACTION_LABELS = {
    "login": "User logged in",
    "logout": "User logged out",
    "view_saved_notes": "Viewed saved notes page",
    "view_saved_notes_page": "Viewed saved notes page",
    "save_notes": "Saved notes",
    "transcription": "Transcription created",
    "ai_summary": "AI summary created",
    "error": "Error occurred",
    "load_stats_fail": "Failed to load statistics",
    "load_activity_fail": "Failed to load activity history",
    "export_note": "Exported note"
}

def _format_timestamp(ts):
    if not ts:
        return "Unknown"
    try:
        return ts.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        try:
            return ts.generation_time.strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            return str(ts)

@app.route("/activity_history")
@login_required
def activity_history():
    try:
        page = request.args.get("page", 1, type=int)
        per_page = 20
        skip = (page - 1) * per_page
        cursor = mongo.db.user_activities.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("timestamp", -1).skip(skip).limit(per_page)
        activities = list(cursor)
        total_activities = mongo.db.user_activities.count_documents(
            {"user_id": ObjectId(current_user.id)}
        )
        has_prev = page > 1
        has_next = skip + per_page < total_activities
        pagination_info = {
            "page": page,
            "per_page": per_page,
            "total": total_activities,
            "has_prev": has_prev,
            "has_next": has_next,
            "prev_page": page - 1 if has_prev else None,
            "next_page": page + 1 if has_next else None,
        }
        for act in activities:
            action = act.get("activity_type") or "unknown"
            act["label"] = ACTION_LABELS.get(action, action.replace("_", " ").title())
            act["details"] = act.get("details", "No details")
            act["timestamp"] = _format_timestamp(act.get("timestamp"))
            metadata = act.get("metadata", {})
            act["provider"] = metadata.get("provider") if isinstance(metadata, dict) else None
        name = current_user.fullname or current_user.username or current_user.email.split("@")[0]
        return render_template(
            "activity_history.html",
            activities=activities,
            pagination=pagination_info,
            pagination_info=pagination_info,
            name=name,
        )
    except Exception as e:
        flash(f"Error loading activity history: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to load activity history: {str(e)}", {"error_type": "activity_history"})
        return redirect(url_for("dashboard"))

@app.route("/regenerate_notes", methods=["POST"])
@login_required
def regenerate_notes():
    try:
        transcript = request.form.get("transcript", "")
        ai_provider = request.form.get("ai_provider", "auto")
        groq_model = request.form.get("groq_model", "llama-3.3-70b-versatile")
        print(f"🔄 Regenerating with provider: {ai_provider}, model: {groq_model}")
        if not transcript:
            flash("No transcript available to regenerate notes!", "danger")
            return redirect(url_for("dashboard"))
        log_user_activity(
            current_user.id,
            "regenerate_notes",
            f"Notes regeneration requested with {ai_provider}",
            {
                "provider": ai_provider,
                "model": groq_model if ai_provider == "groq" else None,
                "transcript_length": len(transcript)
            }
        )
        result = summarize_with_ai(transcript, ai_provider, groq_model)
        if isinstance(result, tuple) and len(result) == 2:
            notes, provider_used = result
            if notes and not notes.startswith("Error"):
                flash(f"Notes regenerated successfully using {provider_used}!", "success")
                update_user_stats(current_user.id, "regenerate_notes")
            else:
                flash(f"AI Processing failed: {notes}", "danger")
                notes = None
                provider_used = None
        else:
            notes = "Error: Failed to process with AI"
            provider_used = "none"
            flash("AI Processing failed: Invalid response", "danger")
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        ai_providers = get_ai_provider_status()
        quick_stats = get_user_statistics(current_user.id, days=7)
        return render_template("dashboard.html",
                             name=name,
                             transcript=transcript,
                             notes=notes,
                             provider_used=provider_used,
                             ai_providers=ai_providers,
                             quick_stats=quick_stats)
    except Exception as e:
        log_user_activity(
            current_user.id,
            "error",
            f"Error regenerating notes: {str(e)}",
            {"error_type": "regeneration"}
        )
        flash(f"Error regenerating notes: {str(e)}", "danger")
        return redirect(url_for("dashboard"))

@app.route("/saved_notes")
@login_required
def saved_notes():
    try:
        notes = list(mongo.db.saved_notes.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("created_at", -1))
        log_user_activity(
            current_user.id,
            "view_saved_notes",
            "Viewed saved notes page",
            {"notes_count": len(notes)}
        )
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        return render_template("saved_notes.html", notes=notes, name=name)
    except Exception as e:
        flash(f"Error loading saved notes: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to load saved notes: {str(e)}", {"error_type": "saved_notes"})
        return redirect(url_for("dashboard"))

@app.route("/view_note/<note_id>")
@login_required
def view_note(note_id):
    try:
        if not ObjectId.is_valid(note_id):
            flash("Invalid note ID!", "danger")
            return redirect(url_for("saved_notes"))
        note = mongo.db.saved_notes.find_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        if not note:
            flash("Note not found!", "danger")
            return redirect(url_for("saved_notes"))
        log_user_activity(
            current_user.id,
            "view_note",
            f"Viewed note: {note.get('title', 'Untitled')}",
            {"note_id": str(note_id)}
        )
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        return render_template("view_note.html", note=note, name=name, export_formats=['docx', 'pdf'])
    except Exception as e:
        flash(f"Error loading note: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to load note: {str(e)}", {"error_type": "view_note", "note_id": str(note_id)})
        return redirect(url_for("saved_notes"))

@app.route("/export_note/<note_id>", methods=["GET"])
@login_required
def export_note(note_id):
    try:
        if not ObjectId.is_valid(note_id):
            flash("Invalid note ID!", "danger")
            return redirect(url_for("saved_notes"))
        
        note = mongo.db.saved_notes.find_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        if not note:
            flash("Note not found!", "danger")
            return redirect(url_for("saved_notes"))
        
        export_format = request.args.get("format", "pdf").lower()
        if export_format not in ['pdf', 'docx']:
            flash("Invalid export format!", "danger")
            return redirect(url_for("view_note", note_id=note_id))
        
        title = note.get('title', 'Untitled_Note').replace(' ', '_').replace('/', '_')
        filename = f"{title}_{datetime.now().strftime('%Y%m%d')}"
        
        log_user_activity(
            current_user.id,
            "export_note",
            f"Exported note as {export_format.upper()}",
            {
                "note_id": str(note_id),
                "format": export_format,
                "title": note.get('title', 'Untitled')
            }
        )
        
        if export_format == 'docx':
            buffer = create_word_document(note)
            filename = f"{filename}.docx"
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:  # PDF
            buffer = create_pdf_document(note)
            filename = f"{filename}.pdf"
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/pdf'
            )
    
    except Exception as e:
        flash(f"Error exporting note: {str(e)}", "danger")
        log_user_activity(
            current_user.id,
            "error",
            f"Failed to export note: {str(e)}",
            {"error_type": f"export_{export_format}", "note_id": str(note_id)}
        )
        return redirect(url_for("view_note", note_id=note_id))

@app.route("/delete_note/<note_id>", methods=["POST"])
@login_required
def delete_note(note_id):
    try:
        if not ObjectId.is_valid(note_id):
            flash("Invalid note ID!", "danger")
            return redirect(url_for("saved_notes"))
        note = mongo.db.saved_notes.find_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        result = mongo.db.saved_notes.delete_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        if result.deleted_count > 0:
            log_user_activity(
                current_user.id,
                "delete_note",
                f"Deleted note: {note.get('title', 'Untitled') if note else 'Unknown'}",
                {"note_id": str(note_id)}
            )
            update_user_stats(current_user.id, "delete_notes")
            flash("Note deleted successfully!", "success")
        else:
            flash("Note not found or already deleted!", "danger")
    except Exception as e:
        flash(f"Error deleting note: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to delete note: {str(e)}", {"error_type": "delete_note", "note_id": str(note_id)})
    return redirect(url_for("saved_notes"))

@app.route("/update_note_title/<note_id>", methods=["POST"])
@login_required
def update_note_title(note_id):
    try:
        if not ObjectId.is_valid(note_id):
            flash("Invalid note ID!", "danger")
            return redirect(url_for("view_note", note_id=note_id))
        new_title = request.form.get("title", "").strip()
        if not new_title:
            flash("Title cannot be empty!", "danger")
            return redirect(url_for("view_note", note_id=note_id))
        old_note = mongo.db.saved_notes.find_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        old_title = old_note.get('title', 'Untitled') if old_note else 'Unknown'
        result = mongo.db.saved_notes.update_one(
            {"_id": ObjectId(note_id), "user_id": ObjectId(current_user.id)},
            {"$set": {"title": new_title, "updated_at": datetime.now().isoformat()}}
        )
        if result.modified_count > 0:
            log_user_activity(
                current_user.id,
                "update_note_title",
                f"Updated note title from '{old_title}' to '{new_title}'",
                {"note_id": str(note_id), "old_title": old_title, "new_title": new_title}
            )
            flash("Note title updated successfully!", "success")
        else:
            flash("Failed to update note title!", "danger")
    except Exception as e:
        flash(f"Error updating note: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to update note title: {str(e)}", {"error_type": "update_title", "note_id": str(note_id)})
    return redirect(url_for("view_note", note_id=note_id))

@app.route("/logout")
@login_required
def logout():
    log_user_activity(current_user.id, "logout", "User logged out")
    logout_user()
    flash("You have logged out!", "info")
    return redirect(url_for("index"))

@app.route("/api/user_stats")
@login_required
def user_stats():
    try:
        days = request.args.get('days', 30, type=int)
        stats = get_user_statistics(current_user.id, days)
        return jsonify(stats)
    except Exception as e:
        log_user_activity(current_user.id, "error", f"Failed to fetch user stats: {str(e)}", {"error_type": "api_user_stats"})
        return jsonify({"error": str(e)}), 500

@app.route("/api/daily_activity")
@login_required
def daily_activity():
    try:
        days = request.args.get('days', 30, type=int)
        start_date = datetime.now() - timedelta(days=days)
        daily_data = list(mongo.db.daily_stats.find({
            "user_id": ObjectId(current_user.id),
            "date": {"$gte": start_date}
        }).sort("date", 1))
        chart_data = [
            {
                "date": day["date"].strftime("%Y-%m-%d") if isinstance(day["date"], datetime) else day["date"],
                "transcription_count": day.get("transcription_count", 0),
                "ai_summary_count": day.get("ai_summary_count", 0),
                "save_notes_count": day.get("save_notes_count", 0),
                "login_count": day.get("login_count", 0)
            } for day in daily_data
        ]
        return jsonify(chart_data)
    except Exception as e:
        log_user_activity(current_user.id, "error", f"Failed to fetch daily activity: {str(e)}", {"error_type": "api_daily_activity"})
        return jsonify({"error": str(e)}), 500

@app.route("/api/activity_timeline")
@login_required
def activity_timeline():
    try:
        limit = request.args.get('limit', 10, type=int)
        activities = list(mongo.db.user_activities.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("timestamp", -1).limit(limit))
        formatted_activities = [
            {
                "timestamp": activity["timestamp"].strftime("%Y-%m-%d %H:%M:%S") if isinstance(activity["timestamp"], datetime) else activity["timestamp"],
                "activity_type": activity["activity_type"],
                "details": activity.get("details", "No details"),
                "metadata": activity.get("metadata", {})
            } for activity in activities
        ]
        return jsonify(formatted_activities)
    except Exception as e:
        log_user_activity(current_user.id, "error", f"Failed to fetch activity timeline: {str(e)}", {"error_type": "api_activity_timeline"})
        return jsonify({"error": str(e)}), 500

@app.route("/api/ai_providers")
@login_required
def api_ai_providers():
    try:
        return jsonify(get_ai_provider_status())
    except Exception as e:
        log_user_activity(current_user.id, "error", f"Failed to fetch AI providers: {str(e)}", {"error_type": "api_ai_providers"})
        return jsonify({"error": str(e)}), 500

@app.route("/api/test_ai_providers", methods=["POST"])
@login_required
def test_ai_providers():
    try:
        test_transcript = request.json.get("transcript", "John: Hello team, let's discuss the Q3 budget.")
        print(f"🧪 Testing AI providers with transcript: {test_transcript[:50]}...")
        log_user_activity(
            current_user.id,
            "test_ai_providers",
            "Tested AI providers functionality",
            {"test_transcript_length": len(test_transcript)}
        )
        results = {}
        providers = get_ai_provider_status()
        if providers['openai']['available']:
            try:
                print("🧪 Testing OpenRouter...")
                summary = summarize_with_openai(test_transcript)
                results['openai'] = {
                    'status': 'success' if not summary.startswith('Error') else 'failed',
                    'response': summary[:200] + "..." if len(summary) > 200 else summary
                }
                print(f"✅ OpenRouter test result: {results['openai']['status']}")
            except Exception as e:
                results['openai'] = {'status': 'failed', 'response': str(e)}
                print(f"❌ OpenRouter test failed: {e}")
        else:
            results['openai'] = {'status': 'unavailable', 'response': 'API key not configured'}
        if providers['groq']['available']:
            try:
                print("🧪 Testing Groq...")
                summary = summarize_with_groq(test_transcript, "llama-3.3-70b-versatile")
                results['groq'] = {
                    'status': 'success' if not summary.startswith('Error') else 'failed',
                    'response': summary[:200] + "..." if len(summary) > 200 else summary
                }
                print(f"✅ Groq test result: {results['groq']['status']}")
            except Exception as e:
                results['groq'] = {'status': 'failed', 'response': str(e)}
                print(f"❌ Groq test failed: {e}")
        else:
            results['groq'] = {'status': 'unavailable', 'response': 'API key not configured'}
        if providers['gemini']['available']:
            try:
                print("🧪 Testing Gemini...")
                summary = summarize_with_gemini(test_transcript)
                results['gemini'] = {
                    'status': 'success' if not summary.startswith('Error') else 'failed',
                    'response': summary[:200] + "..." if len(summary) > 200 else summary
                }
                print(f"✅ Gemini test result: {results['gemini']['status']}")
            except Exception as e:
                results['gemini'] = {'status': 'failed', 'response': str(e)}
                print(f"❌ Gemini test failed: {e}")
        else:
            results['gemini'] = {'status': 'unavailable', 'response': 'API key not configured'}
        print(f"🧪 Test results: {results}")
        return jsonify(results)
    except Exception as e:
        print(f"💥 Test API error: {e}")
        log_user_activity(current_user.id, "error", f"Failed to test AI providers: {str(e)}", {"error_type": "api_test_ai_providers"})
        return jsonify({"error": str(e)}), 500

@app.route("/export_data")
@login_required
def export_data():
    try:
        user_data = {
            "user_info": {
                "email": current_user.email,
                "username": current_user.username,
                "fullname": current_user.fullname,
                "export_date": datetime.now().isoformat()
            },
            "saved_notes": [],
            "activities": [],
            "statistics": get_user_statistics(current_user.id, days=365)
        }
        notes = mongo.db.saved_notes.find({"user_id": ObjectId(current_user.id)})
        for note in notes:
            user_data["saved_notes"].append({
                "title": note.get("title"),
                "transcript": note.get("transcript"),
                "notes": note.get("notes"),
                "ai_provider": note.get("ai_provider"),
                "created_at": note.get("created_at") if isinstance(note.get("created_at"), str) else note.get("created_at").isoformat(),
                "updated_at": note.get("updated_at") if isinstance(note.get("updated_at"), str) else note.get("updated_at").isoformat()
            })
        activities = mongo.db.user_activities.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("timestamp", -1).limit(1000)
        for activity in activities:
            user_data["activities"].append({
                "activity_type": activity.get("activity_type"),
                "details": activity.get("details"),
                "metadata": activity.get("metadata"),
                "timestamp": activity.get("timestamp").strftime("%Y-%m-%d %H:%M:%S") if isinstance(activity.get("timestamp"), datetime) else activity.get("timestamp"),
                "ip_address": activity.get("ip_address")
            })
        log_user_activity(
            current_user.id,
            "export_data",
            "User exported their data",
            {
                "notes_count": len(user_data["saved_notes"]),
                "activities_count": len(user_data["activities"])
            }
        )
        return jsonify(user_data), 200, {
            'Content-Disposition': f'attachment; filename=user_data_{current_user.id}_{datetime.now().strftime("%Y%m%d")}.json',
            'Content-Type': 'application/json'
        }
    except Exception as e:
        flash(f"Error exporting data: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to export data: {str(e)}", {"error_type": "export_data"})
        return redirect(url_for("dashboard"))

@app.route("/clear_activity_history", methods=["POST"])
@login_required
def clear_activity_history():
    try:
        days_to_keep = request.form.get("days_to_keep", 30, type=int)
        cutoff_date = datetime.now() - timedelta(days=days_to_keep)
        result = mongo.db.user_activities.delete_many({
            "user_id": ObjectId(current_user.id),
            "timestamp": {"$lt": cutoff_date}
        })
        log_user_activity(
            current_user.id,
            "clear_history",
            f"Cleared activity history older than {days_to_keep} days",
            {"deleted_count": result.deleted_count, "days_kept": days_to_keep}
        )
        flash(f"Cleared {result.deleted_count} old activity records!", "success")
    except Exception as e:
        flash(f"Error clearing history: {str(e)}", "danger")
        log_user_activity(current_user.id, "error", f"Failed to clear activity history: {str(e)}", {"error_type": "clear_history"})
    return redirect(url_for("activity_history"))

def validate_email(email):
    """Validate email format."""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

# def send_note_via_email(recipient_email, note_data, sender_name, message=""):
#     """Send a note via email with optional attachment."""
#     try:
#         if not EMAIL_USER or not EMAIL_PASSWORD:
#             return False, "Email service not configured. Please contact administrator."
        
#         if not validate_email(recipient_email):
#             return False, "Invalid email address format."
        
#         # Create message
#         msg = MIMEMultipart()
#         msg['From'] = FROM_EMAIL
#         msg['To'] = recipient_email
#         msg['Subject'] = f"Meeting Notes: {note_data.get('title', 'Untitled Note')}"
        
#         # Email body
#         body_text = f"""
# Hello,

# {sender_name} has shared meeting notes with you.

# Title: {note_data.get('title', 'Untitled Note')}
# Created: {note_data.get('created_at', 'Unknown')}
# AI Provider: {note_data.get('ai_provider', 'Unknown')}

# {f"Personal Message: {message}" if message else ""}

# --- TRANSCRIPT ---
# {note_data.get('transcript', 'No transcript available')}

# --- NOTES ---
# {note_data.get('notes', 'No notes available')}

# ---
# This email was sent from the Meeting Notes AI system.
# If you did not expect this email, please ignore it.
#         """
        
#         msg.attach(MIMEText(body_text, 'plain'))
        
#         # Connect to server and send email
#         server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
#         server.starttls()
#         server.login(EMAIL_USER, EMAIL_PASSWORD)
        
#         text = msg.as_string()
#         server.sendmail(FROM_EMAIL, recipient_email, text)
#         server.quit()
        
#         return True, "Email sent successfully!"
        
#     except smtplib.SMTPAuthenticationError:
#         return False, "Email authentication failed. Please check email configuration."
#     except smtplib.SMTPRecipientsRefused:
#         return False, "Recipient email address was refused by the server."
#     except smtplib.SMTPException as e:
#         return False, f"SMTP error occurred: {str(e)}"
#     except Exception as e:
#         return False, f"Failed to send email: {str(e)}"
# /

def send_note_via_email(recipient_email, note_data, sender_name, message=""):
    """Send a note via email without attachment."""
    try:
        if not EMAIL_USER or not EMAIL_PASSWORD:
            return False, "Email service not configured. Please contact administrator."
        
        if not validate_email(recipient_email):
            return False, "Invalid email address format."
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = FROM_EMAIL
        msg['To'] = recipient_email
        msg['Subject'] = f"Meeting Notes: {note_data.get('title', 'Untitled Note')}"
        
        # Email body
        body_text = f"""
Hello,

{sender_name} has shared meeting notes with you.

Title: {note_data.get('title', 'Untitled Note')}
Created: {note_data.get('created_at', 'Unknown')}
AI Provider: {note_data.get('ai_provider', 'Unknown')}

{f"Personal Message: {message}" if message else ""}

--- TRANSCRIPT ---
{note_data.get('transcript', 'No transcript available')}

--- NOTES ---
{note_data.get('notes', 'No notes available')}

---
This email was sent from the Meeting Notes AI system.
If you did not expect this email, please ignore it.
        """
        
        msg.attach(MIMEText(body_text, 'plain'))
        
        # Connect to server and send email
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=30)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        server.sendmail(FROM_EMAIL, recipient_email, msg.as_string())
        server.quit()
        
        return True, "Email sent successfully!"
        
    except smtplib.SMTPAuthenticationError as e:
        app.logger.error(f"SMTP Authentication Error: {str(e)}")
        return False, "Email authentication failed. Please check email configuration."
    except smtplib.SMTPRecipientsRefused as e:
        app.logger.error(f"SMTP Recipients Refused: {str(e)}")
        return False, "Recipient email address was refused by the server."
    except smtplib.SMTPException as e:
        app.logger.error(f"SMTP Error: {str(e)}")
        return False, f"SMTP error occurred: {str(e)}"
    except Exception as e:
        app.logger.error(f"Unexpected Error: {str(e)}")
        return False, f"Failed to send email: {str(e)}"

def send_note_with_attachment(recipient_email, note_data, sender_name, message="", attachment_format="pdf"):
    """Send a note via email with PDF or DOCX attachment."""
    try:
        if not EMAIL_USER or not EMAIL_PASSWORD:
            return False, "Email service not configured. Please contact administrator."
        
        if not validate_email(recipient_email):
            return False, "Invalid email address format."
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = FROM_EMAIL
        msg['To'] = recipient_email
        msg['Subject'] = f"Meeting Notes: {note_data.get('title', 'Untitled Note')}"
        
        # Email body
        body_text = f"""
Hello,

{sender_name} has shared meeting notes with you.

Title: {note_data.get('title', 'Untitled Note')}
Created: {note_data.get('created_at', 'Unknown')}
AI Provider: {note_data.get('ai_provider', 'Unknown')}

{f"Personal Message: {message}" if message else ""}

Please find the detailed notes attached to this email.

---
This email was sent from the Meeting Notes AI system.
If you did not expect this email, please ignore it.
        """
        
        msg.attach(MIMEText(body_text, 'plain'))
        
        # Create attachment
        if attachment_format.lower() == 'pdf':
            attachment_buffer = create_pdf_document(note_data)
            filename = f"{note_data.get('title', 'Meeting_Notes').replace(' ', '_')}.pdf"
            mimetype = 'application/pdf'
        else:  # docx
            attachment_buffer = create_word_document(note_data)
            filename = f"{note_data.get('title', 'Meeting_Notes').replace(' ', '_')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        attachment = MIMEApplication(attachment_buffer.read(), _subtype=attachment_format)
        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
        msg.attach(attachment)
        
        # Connect to server and send email
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=30)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        server.sendmail(FROM_EMAIL, recipient_email, msg.as_string())
        server.quit()
        
        return True, "Email with attachment sent successfully!"
        
    except smtplib.SMTPAuthenticationError as e:
        app.logger.error(f"SMTP Authentication Error: {str(e)}")
        return False, "Email authentication failed. Please check email configuration."
    except smtplib.SMTPRecipientsRefused as e:
        app.logger.error(f"SMTP Recipients Refused: {str(e)}")
        return False, "Recipient email address was refused by the server."
    except smtplib.SMTPException as e:
        app.logger.error(f"SMTP Error: {str(e)}")
        return False, f"SMTP error occurred: {str(e)}"
    except Exception as e:
        app.logger.error(f"Unexpected Error: {str(e)}")
        return False, f"Failed to send email: {str(e)}"
def send_note_with_attachment(recipient_email, note_data, sender_name, message="", attachment_format="pdf"):
    """Send a note via email with PDF or DOCX attachment."""
    try:
        if not EMAIL_USER or not EMAIL_PASSWORD:
            return False, "Email service not configured. Please contact administrator."
        
        if not validate_email(recipient_email):
            return False, "Invalid email address format."
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = FROM_EMAIL
        msg['To'] = recipient_email
        msg['Subject'] = f"Meeting Notes: {note_data.get('title', 'Untitled Note')}"
        
        # Email body
        body_text = f"""
Hello,

{sender_name} has shared meeting notes with you.

Title: {note_data.get('title', 'Untitled Note')}
Created: {note_data.get('created_at', 'Unknown')}
AI Provider: {note_data.get('ai_provider', 'Unknown')}

{f"Personal Message: {message}" if message else ""}

Please find the detailed notes attached to this email.

---
This email was sent from the Meeting Notes AI system.
If you did not expect this email, please ignore it.
        """
        
        msg.attach(MIMEText(body_text, 'plain'))
        
        # Create attachment
        if attachment_format.lower() == 'pdf':
            attachment_buffer = create_pdf_document(note_data)
            filename = f"{note_data.get('title', 'Meeting_Notes').replace(' ', '_')}.pdf"
            mimetype = 'application/pdf'
        else:  # docx
            attachment_buffer = create_word_document(note_data)
            filename = f"{note_data.get('title', 'Meeting_Notes').replace(' ', '_')}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        attachment = MIMEApplication(attachment_buffer.read(), _subtype=attachment_format)
        attachment.add_header('Content-Disposition', 'attachment', filename=filename)
        msg.attach(attachment)
        
        # Connect to server and send email
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        
        text = msg.as_string()
        server.sendmail(FROM_EMAIL, recipient_email, text)
        server.quit()
        
        return True, "Email with attachment sent successfully!"
        
    except Exception as e:
        return False, f"Failed to send email: {str(e)}"

# Add this route to your Flask app
# @app.route("/email_note/<note_id>", methods=["GET", "POST"])
# @login_required
# def email_note(note_id):
#     try:
#         if not ObjectId.is_valid(note_id):
#             flash("Invalid note ID!", "danger")
#             return redirect(url_for("saved_notes"))
        
#         note = mongo.db.saved_notes.find_one({
#             "_id": ObjectId(note_id),
#             "user_id": ObjectId(current_user.id)
#         })
        
#         if not note:
#             flash("Note not found!", "danger")
#             return redirect(url_for("saved_notes"))
        
#         if request.method == "POST":
#             recipient_email = request.form.get("recipient_email", "").strip()
#             personal_message = request.form.get("personal_message", "").strip()
#             include_attachment = request.form.get("include_attachment") == "on"
#             attachment_format = request.form.get("attachment_format", "pdf")
            
#             if not recipient_email:
#                 flash("Recipient email is required!", "danger")
#                 return redirect(url_for("email_note", note_id=note_id))
            
#             if not validate_email(recipient_email):
#                 flash("Please enter a valid email address!", "danger")
#                 return redirect(url_for("email_note", note_id=note_id))
            
#             # Check email configuration
#             if not EMAIL_USER or not EMAIL_PASSWORD:
#                 flash("Email service is not configured. Please contact administrator.", "danger")
#                 return redirect(url_for("email_note", note_id=note_id))
            
#             sender_name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
            
#             # Send email
#             if include_attachment:
#                 success, message = send_note_with_attachment(
#                     recipient_email, 
#                     note, 
#                     sender_name, 
#                     personal_message,
#                     attachment_format
#                 )
#             else:
#                 success, message = send_note_via_email(
#                     recipient_email, 
#                     note, 
#                     sender_name, 
#                     personal_message
#                 )
            
#             if success:
#                 # Log the activity
#                 log_user_activity(
#                     current_user.id,
#                     "email_note",
#                     f"Emailed note '{note.get('title', 'Untitled')}' to {recipient_email}",
#                     {
#                         "note_id": str(note_id),
#                         "recipient_email": recipient_email,
#                         "with_attachment": include_attachment,
#                         "attachment_format": attachment_format if include_attachment else None,
#                         "message_length": len(personal_message)
#                     }
#                 )
#                 update_user_stats(current_user.id, "email_notes")
#                 flash(message, "success")
#                 return redirect(url_for("view_note", note_id=note_id))
#             else:
#                 flash(message, "danger")
#                 return redirect(url_for("email_note", note_id=note_id))
        
#         # GET request - show email form
#         name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
#         email_configured = bool(EMAIL_USER and EMAIL_PASSWORD)
        
#         return render_template("email_note.html", 
#                              note=note, 
#                              name=name,
#                              email_configured=email_configured)
        
#     except Exception as e:
#         flash(f"Error processing email request: {str(e)}", "danger")
#         log_user_activity(
#             current_user.id, 
#             "error", 
#             f"Failed to email note: {str(e)}", 
#             {"error_type": "email_note", "note_id": str(note_id)}
#         )
#         return redirect(url_for("view_note", note_id=note_id))
# /
@app.route("/email_note/<note_id>", methods=["GET", "POST"])
@login_required
def email_note(note_id):
    try:
        if not ObjectId.is_valid(note_id):
            flash("Invalid note ID!", "danger")
            return redirect(url_for("saved_notes"))
        
        note = mongo.db.saved_notes.find_one({
            "_id": ObjectId(note_id),
            "user_id": ObjectId(current_user.id)
        })
        
        if not note:
            flash("Note not found!", "danger")
            return redirect(url_for("saved_notes"))
        
        if request.method == "POST":
            recipient_email = request.form.get("recipient_email", "").strip()
            personal_message = request.form.get("personal_message", "").strip()
            include_attachment = request.form.get("include_attachment") == "on"
            attachment_format = request.form.get("attachment_format", "pdf").lower()
            
            if not recipient_email:
                flash("Recipient email is required!", "danger")
                return redirect(url_for("email_note", note_id=note_id))
            
            if not validate_email(recipient_email):
                flash("Please enter a valid email address!", "danger")
                return redirect(url_for("email_note", note_id=note_id))
            
            # Check email configuration
            if not EMAIL_USER or not EMAIL_PASSWORD:
                flash("Email service is not configured. Please contact administrator.", "danger")
                return redirect(url_for("email_note", note_id=note_id))
            
            sender_name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
            
            # Send email
            if include_attachment:
                success, message = send_note_with_attachment(
                    recipient_email, 
                    note, 
                    sender_name, 
                    personal_message,
                    attachment_format
                )
            else:
                success, message = send_note_via_email(
                    recipient_email, 
                    note, 
                    sender_name, 
                    personal_message
                )
            
            if success:
                # Log the activity
                log_user_activity(
                    current_user.id,
                    "email_note",
                    f"Emailed note '{note.get('title', 'Untitled')}' to {recipient_email}",
                    {
                        "note_id": str(note_id),
                        "recipient_email": recipient_email,
                        "with_attachment": include_attachment,
                        "attachment_format": attachment_format if include_attachment else None,
                        "message_length": len(personal_message)
                    }
                )
                update_user_stats(current_user.id, "email_notes")
                flash(message, "success")
                return redirect(url_for("saved_notes"))  # Redirect to saved_notes for consistency
            else:
                flash(message, "danger")
                return redirect(url_for("email_note", note_id=note_id))
        
        # GET request - show email form
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        email_configured = bool(EMAIL_USER and EMAIL_PASSWORD)
        
        return render_template("email_note.html", 
                             note=note, 
                             name=name,
                             email_configured=email_configured)
        
    except Exception as e:
        app.logger.error(f"Error processing email request: {str(e)}")
        flash(f"Error processing email request: {str(e)}", "danger")
        log_user_activity(
            current_user.id, 
            "error", 
            f"Failed to email note: {str(e)}", 
            {"error_type": "email_note", "note_id": str(note_id)}
        )
        return redirect(url_for("saved_notes"))
# Add this route for bulk email functionality
@app.route("/bulk_email_notes", methods=["GET", "POST"])
@login_required
def bulk_email_notes():
    try:
        if request.method == "POST":
            selected_notes = request.form.getlist("selected_notes")
            recipient_emails = request.form.get("recipient_emails", "").strip()
            personal_message = request.form.get("personal_message", "").strip()
            include_attachment = request.form.get("include_attachment") == "on"
            attachment_format = request.form.get("attachment_format", "pdf")
            
            if not selected_notes:
                flash("Please select at least one note to send!", "danger")
                return redirect(url_for("bulk_email_notes"))
            
            if not recipient_emails:
                flash("Please enter at least one recipient email!", "danger")
                return redirect(url_for("bulk_email_notes"))
            
            # Parse email addresses (comma or newline separated)
            emails = [email.strip() for email in re.split('[,\n]', recipient_emails) if email.strip()]
            invalid_emails = [email for email in emails if not validate_email(email)]
            
            if invalid_emails:
                flash(f"Invalid email addresses: {', '.join(invalid_emails)}", "danger")
                return redirect(url_for("bulk_email_notes"))
            
            # Check email configuration
            if not EMAIL_USER or not EMAIL_PASSWORD:
                flash("Email service is not configured. Please contact administrator.", "danger")
                return redirect(url_for("bulk_email_notes"))
            
            sender_name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
            success_count = 0
            failure_count = 0
            
            # Send emails
            for note_id in selected_notes:
                if not ObjectId.is_valid(note_id):
                    continue
                    
                note = mongo.db.saved_notes.find_one({
                    "_id": ObjectId(note_id),
                    "user_id": ObjectId(current_user.id)
                })
                
                if not note:
                    continue
                
                for email in emails:
                    if include_attachment:
                        success, message = send_note_with_attachment(
                            email, note, sender_name, personal_message, attachment_format
                        )
                    else:
                        success, message = send_note_via_email(
                            email, note, sender_name, personal_message
                        )
                    
                    if success:
                        success_count += 1
                    else:
                        failure_count += 1
            
            # Log bulk email activity
            log_user_activity(
                current_user.id,
                "bulk_email_notes",
                f"Bulk emailed {len(selected_notes)} notes to {len(emails)} recipients",
                {
                    "notes_count": len(selected_notes),
                    "recipients_count": len(emails),
                    "success_count": success_count,
                    "failure_count": failure_count,
                    "with_attachment": include_attachment
                }
            )
            
            if success_count > 0:
                update_user_stats(current_user.id, "bulk_email_notes", success_count)
                flash(f"Successfully sent {success_count} emails!" + 
                     (f" {failure_count} failed." if failure_count > 0 else ""), 
                     "success" if failure_count == 0 else "warning")
            else:
                flash("Failed to send any emails. Please check your configuration.", "danger")
            
            return redirect(url_for("saved_notes"))
        
        # GET request - show bulk email form
        notes = list(mongo.db.saved_notes.find(
            {"user_id": ObjectId(current_user.id)}
        ).sort("created_at", -1))
        
        name = current_user.fullname or current_user.username or current_user.email.split('@')[0]
        email_configured = bool(EMAIL_USER and EMAIL_PASSWORD)
        
        return render_template("bulk_email_notes.html", 
                             notes=notes, 
                             name=name,
                             email_configured=email_configured)
        
    except Exception as e:
        flash(f"Error processing bulk email request: {str(e)}", "danger")
        log_user_activity(
            current_user.id, 
            "error", 
            f"Failed bulk email operation: {str(e)}", 
            {"error_type": "bulk_email_notes"}
        )
        return redirect(url_for("saved_notes"))
if __name__ == "__main__":
    app.run(debug=True)